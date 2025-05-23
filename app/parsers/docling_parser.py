import os
import base64
from typing import Dict, List, Any, Optional
from PIL import Image
import numpy as np
import pandas as pd
from docx import Document
import json
import io

try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("Warning: Docling not available. Install with: pip install docling")

class DoclingParser:
    """Advanced document parser using Docling for superior multimodal extraction"""
    
    def __init__(self):
        if not DOCLING_AVAILABLE:
            raise ImportError("Docling is required. Install with: pip install docling")
        
        # Initialize basic Docling converter
        try:
            self.converter = DocumentConverter()
        except Exception as e:
            raise ImportError(f"Failed to initialize DocumentConverter: {e}")
        
        self.supported_formats = {
            'pdf': self._parse_with_docling,
            'docx': self._parse_docx,
            'doc': self._parse_doc_fallback,
            'xlsx': self._parse_excel,
            'xls': self._parse_excel,
            'csv': self._parse_csv,
            'png': self._parse_image_with_docling,
            'jpg': self._parse_image_with_docling,
            'jpeg': self._parse_image_with_docling,
            'tiff': self._parse_image_with_docling,
            'bmp': self._parse_image_with_docling
        }
    
    def parse_document(self, file_path: str) -> Dict:
        """Parse any supported document format using Docling"""
        if not os.path.exists(file_path):
            return {'error': f'File not found: {file_path}'}
        
        file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        
        if file_ext not in self.supported_formats:
            return {'error': f'Unsupported file format: {file_ext}'}
        
        try:
            return self.supported_formats[file_ext](file_path)
        except Exception as e:
            return {'error': f'Error parsing {file_ext} file: {str(e)}'}
    
    def _parse_with_docling(self, file_path: str) -> Dict:
        """Parse document using Docling's advanced capabilities"""
        try:
            # Convert document
            result = self.converter.convert(file_path)
            doc = result.document
            
            parsed_result = {
                'text_content': '',
                'tables': [],
                'images': [],
                'charts': [],
                'figures': [],
                'metadata': {
                    'title': getattr(doc, 'title', ''),
                    'pages': len(getattr(doc, 'pages', [])),
                    'format': 'pdf'
                },
                'structure': {
                    'headings': [],
                    'sections': [],
                    'lists': []
                }
            }
            
            # Extract text content
            if hasattr(doc, 'export_to_markdown'):
                parsed_result['text_content'] = doc.export_to_markdown()
            elif hasattr(doc, 'text'):
                parsed_result['text_content'] = doc.text
            
            # Extract structured elements
            if hasattr(doc, 'iterate_items'):
                for item in doc.iterate_items():
                    self._process_docling_item(item, parsed_result)
            
            # Extract tables
            if hasattr(doc, 'tables'):
                for i, table in enumerate(doc.tables):
                    table_data = self._extract_table_from_docling(table)
                    if table_data:
                        parsed_result['tables'].append({
                            'table_index': i,
                            'data': table_data,
                            'analysis': self._analyze_table(table_data),
                            'bbox': getattr(table, 'bbox', None),
                            'page': getattr(table, 'page', None)
                        })
            
            # Extract figures and images
            if hasattr(doc, 'pictures'):
                for i, picture in enumerate(doc.pictures):
                    image_info = self._extract_image_from_docling(picture, i)
                    if image_info:
                        parsed_result['images'].append(image_info)
                        
                        # Check if image contains charts
                        if self._is_chart_or_graph(image_info.get('analysis', {})):
                            parsed_result['charts'].append({
                                'chart_index': len(parsed_result['charts']),
                                'source_image': i,
                                'type': image_info.get('analysis', {}).get('chart_type', 'unknown'),
                                'extracted_data': image_info.get('analysis', {}).get('data_points', []),
                                'analysis': image_info.get('analysis', {})
                            })
            
            # Extract document structure
            if hasattr(doc, 'headings'):
                parsed_result['structure']['headings'] = [
                    {
                        'text': h.text,
                        'level': getattr(h, 'level', 1),
                        'page': getattr(h, 'page', None)
                    } for h in doc.headings
                ]
            
            return parsed_result
            
        except Exception as e:
            return {'error': f'Docling parsing failed: {str(e)}'}
    
    def _process_docling_item(self, item, parsed_result: Dict):
        """Process individual items from Docling document structure"""
        try:
            item_type = getattr(item, 'type', 'unknown')
            
            if item_type == 'table':
                # Already handled in main table extraction
                pass
            elif item_type == 'figure':
                figure_info = {
                    'type': 'figure',
                    'caption': getattr(item, 'caption', ''),
                    'bbox': getattr(item, 'bbox', None),
                    'page': getattr(item, 'page', None)
                }
                parsed_result['figures'].append(figure_info)
            elif item_type == 'list':
                list_items = getattr(item, 'items', [])
                if list_items:
                    parsed_result['structure']['lists'].append({
                        'items': [str(list_item) for list_item in list_items],
                        'page': getattr(item, 'page', None)
                    })
            elif item_type == 'section':
                section_info = {
                    'title': getattr(item, 'title', ''),
                    'content': getattr(item, 'text', ''),
                    'page': getattr(item, 'page', None)
                }
                parsed_result['structure']['sections'].append(section_info)
                
        except Exception as e:
            # Log error but continue processing
            print(f"Warning: Error processing Docling item: {e}")
    
    def _extract_table_from_docling(self, table) -> List[List[str]]:
        """Extract table data from Docling table object"""
        try:
            # Handle different types of table objects
            if hasattr(table, 'data') and table.data is not None:
                # Check if it's a TableData object with table property
                if hasattr(table.data, 'table'):
                    table_data = table.data.table
                    # Convert to list of lists
                    result = []
                    for row in table_data:
                        row_data = []
                        for cell in row:
                            cell_text = str(cell) if cell is not None else ""
                            row_data.append(cell_text)
                        result.append(row_data)
                    return result
                
                # Try to convert TableData to list directly
                elif hasattr(table.data, '__iter__'):
                    result = []
                    for row in table.data:
                        if hasattr(row, '__iter__') and not isinstance(row, str):
                            row_data = [str(cell) for cell in row]
                        else:
                            row_data = [str(row)]
                        result.append(row_data)
                    return result
                else:
                    return [[str(table.data)]]
            
            elif hasattr(table, 'to_dataframe'):
                df = table.to_dataframe()
                return [df.columns.tolist()] + df.values.tolist()
            
            elif hasattr(table, 'cells'):
                # Manual cell extraction
                rows = {}
                for cell in table.cells:
                    row_idx = getattr(cell, 'row', 0)
                    col_idx = getattr(cell, 'col', 0)
                    text = getattr(cell, 'text', '')
                    
                    if row_idx not in rows:
                        rows[row_idx] = {}
                    rows[row_idx][col_idx] = text
                
                # Convert to list format
                table_data = []
                for row_idx in sorted(rows.keys()):
                    row_data = []
                    max_col = max(rows[row_idx].keys()) if rows[row_idx] else 0
                    for col_idx in range(max_col + 1):
                        row_data.append(rows[row_idx].get(col_idx, ''))
                    table_data.append(row_data)
                
                return table_data
            
            else:
                # Try to extract text representation
                text_repr = str(table)
                if text_repr and text_repr != str(type(table)):
                    return [[text_repr]]
                return []
                
        except Exception as e:
            print(f"Warning: Error extracting table: {e}")
            return []
    
    def _extract_image_from_docling(self, picture, index: int) -> Optional[Dict]:
        """Extract image information from Docling picture object"""
        try:
            image_info = {
                'image_index': index,
                'bbox': getattr(picture, 'bbox', None),
                'page': getattr(picture, 'page', None),
                'caption': getattr(picture, 'caption', ''),
                'analysis': {}
            }
            
            # Get image data if available
            if hasattr(picture, 'image'):
                image_data = picture.image
                if image_data:
                    # Analyze image content
                    image_info['analysis'] = self._analyze_image_content(image_data)
                    
                    # Convert to base64 for storage
                    if isinstance(image_data, bytes):
                        image_info['base64'] = base64.b64encode(image_data).decode()
                    elif hasattr(image_data, 'tobytes'):
                        image_info['base64'] = base64.b64encode(image_data.tobytes()).decode()
            
            return image_info
            
        except Exception as e:
            print(f"Warning: Error extracting image {index}: {e}")
            return None
    
    def _parse_image_with_docling(self, file_path: str) -> Dict:
        """Parse standalone image files using Docling"""
        try:
            # For image files, we can still use Docling if it supports them
            # Otherwise fall back to basic image processing
            result = self.converter.convert(file_path)
            
            if result and hasattr(result, 'document'):
                return self._parse_with_docling(file_path)
            else:
                # Fallback to basic image processing
                return self._parse_image_basic(file_path)
                
        except Exception:
            # Fallback to basic image processing
            return self._parse_image_basic(file_path)
    
    def _parse_image_basic(self, file_path: str) -> Dict:
        """Basic image parsing without Docling"""
        try:
            with open(file_path, 'rb') as f:
                image_data = f.read()
            
            image = Image.open(io.BytesIO(image_data))
            
            result = {
                'text_content': '',
                'images': [{
                    'image_index': 0,
                    'analysis': self._analyze_image_content(image_data),
                    'dimensions': image.size,
                    'format': image.format
                }],
                'metadata': {
                    'format': 'image',
                    'dimensions': image.size
                }
            }
            
            # Extract any text content from analysis
            analysis = result['images'][0]['analysis']
            if 'ocr_text' in analysis:
                result['text_content'] = analysis['ocr_text']
            
            # Check for charts
            if self._is_chart_or_graph(analysis):
                result['charts'] = [{
                    'chart_index': 0,
                    'type': analysis.get('chart_type', 'unknown'),
                    'extracted_data': analysis.get('data_points', []),
                    'analysis': analysis
                }]
            
            return result
            
        except Exception as e:
            return {'error': f'Basic image parsing failed: {str(e)}'}
    
    def _parse_docx(self, file_path: str) -> Dict:
        """Parse Word document (keeping existing implementation)"""
        doc = Document(file_path)
        
        result = {
            'text_content': '',
            'images': [],
            'tables': [],
            'metadata': {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or ''
            }
        }
        
        # Extract text
        for paragraph in doc.paragraphs:
            result['text_content'] += paragraph.text + '\n'
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            result['tables'].append({
                'table_index': table_idx,
                'data': table_data,
                'analysis': self._analyze_table(table_data)
            })
        
        return result
    
    def _parse_doc_fallback(self, file_path: str) -> Dict:
        """Fallback for legacy .doc files"""
        return {'error': 'Legacy .doc format requires conversion to .docx'}
    
    def _parse_excel(self, file_path: str) -> Dict:
        """Parse Excel file (keeping existing implementation)"""
        result = {
            'text_content': '',
            'tables': [],
            'charts': [],
            'metadata': {}
        }
        
        try:
            excel_file = pd.ExcelFile(file_path)
            result['metadata']['sheets'] = excel_file.sheet_names
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                table_data = [df.columns.tolist()] + df.values.tolist()
                
                result['tables'].append({
                    'sheet_name': sheet_name,
                    'data': table_data,
                    'analysis': self._analyze_table(table_data),
                    'statistics': {
                        'rows': len(df),
                        'columns': len(df.columns),
                        'numeric_columns': len(df.select_dtypes(include=[np.number]).columns)
                    }
                })
                
                result['text_content'] += f"\n--- Sheet: {sheet_name} ---\n"
                result['text_content'] += f"Columns: {', '.join(df.columns)}\n"
                result['text_content'] += f"Shape: {df.shape[0]} rows x {df.shape[1]} columns\n"
        
        except Exception as e:
            result['error'] = f"Error parsing Excel file: {str(e)}"
        
        return result
    
    def _parse_csv(self, file_path: str) -> Dict:
        """Parse CSV file (keeping existing implementation)"""
        try:
            df = pd.read_csv(file_path)
            table_data = [df.columns.tolist()] + df.values.tolist()
            
            return {
                'text_content': f"CSV file with {df.shape[0]} rows and {df.shape[1]} columns\nColumns: {', '.join(df.columns)}",
                'tables': [{
                    'data': table_data,
                    'analysis': self._analyze_table(table_data),
                    'statistics': {
                        'rows': len(df),
                        'columns': len(df.columns),
                        'numeric_columns': len(df.select_dtypes(include=[np.number]).columns)
                    }
                }],
                'metadata': {'format': 'csv'}
            }
        except Exception as e:
            return {'error': f"Error parsing CSV: {str(e)}"}
    
    def _analyze_table(self, table_data: List[List]) -> Dict:
        """Analyze table content for policy-relevant information"""
        if not table_data or len(table_data) < 2:
            return {'type': 'empty_table'}
        
        headers = table_data[0] if table_data else []
        data_rows = table_data[1:] if len(table_data) > 1 else []
        
        analysis = {
            'headers': headers,
            'row_count': len(data_rows),
            'column_count': len(headers),
            'data_types': [],
            'potential_metrics': [],
            'compliance_indicators': []
        }
        
        # Analyze headers for policy-relevant content
        policy_keywords = ['compliance', 'threshold', 'limit', 'requirement', 'standard', 'benchmark', 'target', 'kpi', 'metric']
        for header in headers:
            if any(keyword in str(header).lower() for keyword in policy_keywords):
                analysis['compliance_indicators'].append(header)
        
        # Look for numeric data that could be thresholds or metrics
        for col_idx, header in enumerate(headers):
            if col_idx < len(data_rows[0]) if data_rows else False:
                sample_values = [row[col_idx] for row in data_rows[:5] if col_idx < len(row)]
                numeric_count = sum(1 for val in sample_values if self._is_numeric(val))
                
                if numeric_count > len(sample_values) * 0.7:  # 70% numeric
                    analysis['potential_metrics'].append({
                        'column': header,
                        'type': 'numeric',
                        'sample_values': sample_values
                    })
        
        return analysis
    
    def _analyze_image_content(self, image_data: bytes) -> Dict:
        """Analyze image content - simplified without OCR dependencies"""
        try:
            image = Image.open(io.BytesIO(image_data))
            
            analysis = {
                'dimensions': image.size,
                'format': image.format,
                'mode': image.mode,
                'has_transparency': 'transparency' in image.info,
                'is_chart': False,
                'chart_type': 'unknown',
                'data_points': []
            }
            
            # Simple heuristics for chart detection
            width, height = image.size
            aspect_ratio = width / height
            
            # Charts often have specific aspect ratios
            if 0.5 <= aspect_ratio <= 2.0 and min(width, height) > 200:
                analysis['is_chart'] = True
                
                # Simple chart type guessing based on image properties
                if aspect_ratio > 1.5:
                    analysis['chart_type'] = 'bar_chart'
                elif 0.8 <= aspect_ratio <= 1.2:
                    analysis['chart_type'] = 'pie_chart'
                else:
                    analysis['chart_type'] = 'line_chart'
            
            return analysis
            
        except Exception as e:
            return {'error': f'Image analysis failed: {str(e)}'}
    
    def _is_chart_or_graph(self, analysis: Dict) -> bool:
        """Determine if image contains charts or graphs"""
        return analysis.get('is_chart', False) or analysis.get('chart_type', 'unknown') != 'unknown'
    
    def _is_numeric(self, value) -> bool:
        """Check if value is numeric"""
        try:
            float(str(value))
            return True
        except (ValueError, TypeError):
            return False
    
    def get_document_summary(self, parsed_document: Dict) -> Dict:
        """Generate a summary of parsed document content"""
        summary = {
            'content_types': [],
            'policy_indicators': [],
            'metrics_found': [],
            'tables_count': len(parsed_document.get('tables', [])),
            'images_count': len(parsed_document.get('images', [])),
            'charts_count': len(parsed_document.get('charts', [])),
            'figures_count': len(parsed_document.get('figures', [])),
            'text_length': len(parsed_document.get('text_content', '')),
            'structure_elements': len(parsed_document.get('structure', {}).get('headings', []))
        }
        
        # Analyze content types
        if summary['text_length'] > 0:
            summary['content_types'].append('text')
        if summary['tables_count'] > 0:
            summary['content_types'].append('tables')
        if summary['images_count'] > 0:
            summary['content_types'].append('images')
        if summary['charts_count'] > 0:
            summary['content_types'].append('charts')
        if summary['figures_count'] > 0:
            summary['content_types'].append('figures')
        if summary['structure_elements'] > 0:
            summary['content_types'].append('structured_document')
        
        # Look for policy-related keywords
        text = parsed_document.get('text_content', '').lower()
        policy_keywords = [
            'policy', 'compliance', 'requirement', 'standard', 'guideline',
            'threshold', 'limit', 'benchmark', 'target', 'kpi', 'metric',
            'audit', 'assessment', 'evaluation', 'criteria', 'procedure'
        ]
        
        for keyword in policy_keywords:
            if keyword in text:
                summary['policy_indicators'].append(keyword)
        
        # Extract potential metrics from tables
        for table in parsed_document.get('tables', []):
            metrics = table.get('analysis', {}).get('potential_metrics', [])
            summary['metrics_found'].extend([m['column'] for m in metrics])
        
        return summary